"""
Funções utilitárias do Nutri IA:
- chamada ao Gemini
- cálculos nutricionais (IMC, TMB, GET, faixa calórica)
- geração de arquivos .docx e .pdf para entrega ao paciente
"""

import io
import re
from datetime import datetime

from google import genai
from google.genai import types as genai_types

from config import (
    GEMINI_MODEL,
    FAIXAS_IMC,
    FATORES_ATIVIDADE,
    AJUSTE_OBJETIVO,
)


# ---------------------------------------------------------------------------
# Gemini
# ---------------------------------------------------------------------------

def get_gemini_response(
    prompt: str,
    api_key: str,
    system_instruction: str,
    buscar_na_web: bool = False,
) -> str:
    """Envia um prompt ao Gemini com a instrução de sistema do módulo atual.

    Quando buscar_na_web=True, ativa o Grounding with Google Search: o Gemini
    pesquisa na web antes de responder e a função anexa ao final do texto uma
    lista "Fontes consultadas" com os links usados. Se a chave/conta não tiver
    esse recurso liberado (ele pode exigir faturamento habilitado no Google
    Cloud, dependendo do tipo de chave), a função tenta de novo sem busca em
    vez de quebrar — e sinaliza isso no texto retornado.

    Levanta uma exceção com mensagem amigável em português caso a chamada falhe
    (chave inválida, limite de uso atingido, etc.) para que a interface possa
    exibir o erro de forma clara ao nutricionista.
    """
    if not api_key:
        raise ValueError("Cole sua chave da API do Gemini na barra lateral antes de continuar.")

    client = genai.Client(api_key=api_key)

    if buscar_na_web:
        try:
            return _gerar_com_busca(client, prompt, system_instruction)
        except Exception as exc:  # noqa: BLE001 - fallback deliberado, sem busca
            aviso = (
                "\n\n---\n_⚠️ Busca em fontes na web não disponível com esta chave de API "
                "(pode exigir faturamento habilitado no Google Cloud). Resposta gerada com "
                "o conhecimento do modelo, sem consulta em tempo real._"
            )
            texto = _gerar_sem_busca(client, prompt, system_instruction)
            return texto + aviso

    return _gerar_sem_busca(client, prompt, system_instruction)


def _gerar_com_busca(client, prompt: str, system_instruction: str) -> str:
    """Chama o Gemini com Grounding with Google Search ativado e anexa as
    fontes usadas ao final da resposta."""
    grounding_tool = genai_types.Tool(google_search=genai_types.GoogleSearch())
    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt,
        config=genai_types.GenerateContentConfig(
            system_instruction=system_instruction,
            tools=[grounding_tool],
        ),
    )
    texto = response.text

    try:
        metadata = response.candidates[0].grounding_metadata
        chunks = metadata.grounding_chunks if metadata else None
    except (AttributeError, IndexError):
        chunks = None

    if chunks:
        fontes = []
        vistos = set()
        for chunk in chunks:
            web = getattr(chunk, "web", None)
            if not web or not web.uri:
                continue
            if web.uri in vistos:
                continue
            vistos.add(web.uri)
            titulo = web.title or web.uri
            fontes.append(f"- [{titulo}]({web.uri})")
        if fontes:
            texto += "\n\n---\n**Fontes consultadas:**\n" + "\n".join(fontes)

    return texto


def _gerar_sem_busca(client, prompt: str, system_instruction: str) -> str:
    try:
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                system_instruction=system_instruction,
            ),
        )
        return response.text
    except Exception as exc:  # noqa: BLE001 - queremos capturar qualquer erro da API
        mensagem = str(exc)
        if "429" in mensagem or "quota" in mensagem.lower():
            raise RuntimeError(
                "Limite de uso da API do Gemini atingido. Aguarde um pouco antes de tentar "
                "novamente ou verifique sua cota em aistudio.google.com."
            ) from exc
        if "API key" in mensagem or "API_KEY_INVALID" in mensagem:
            raise RuntimeError(
                "Chave da API inválida. Confira a chave colada na barra lateral."
            ) from exc
        raise RuntimeError(f"Erro ao consultar o Gemini: {mensagem}") from exc


# ---------------------------------------------------------------------------
# Calculadora nutricional
# ---------------------------------------------------------------------------

def calcular_imc(peso_kg: float, altura_cm: float) -> dict:
    altura_m = altura_cm / 100
    imc = peso_kg / (altura_m ** 2)
    classificacao = next(
        label for minimo, maximo, label in FAIXAS_IMC if minimo <= imc < maximo
    )
    return {"imc": round(imc, 1), "classificacao": classificacao}


def calcular_tmb(
    peso_kg: float,
    altura_cm: float,
    idade: int,
    sexo: str,
    formula: str = "Mifflin-St Jeor",
) -> float:
    """Calcula a Taxa Metabólica Basal (TMB) em kcal/dia.

    sexo: "Feminino" ou "Masculino"
    formula: "Mifflin-St Jeor" (padrão) ou "Harris-Benedict"
    """
    if formula == "Mifflin-St Jeor":
        base = 10 * peso_kg + 6.25 * altura_cm - 5 * idade
        tmb = base + 5 if sexo == "Masculino" else base - 161
    else:  # Harris-Benedict revisada
        if sexo == "Masculino":
            tmb = 88.362 + (13.397 * peso_kg) + (4.799 * altura_cm) - (5.677 * idade)
        else:
            tmb = 447.593 + (9.247 * peso_kg) + (3.098 * altura_cm) - (4.330 * idade)
    return round(tmb, 0)


def calcular_get(tmb: float, nivel_atividade: str) -> float:
    fator = FATORES_ATIVIDADE[nivel_atividade]
    return round(tmb * fator, 0)


def calcular_faixa_calorica(get: float, objetivo: str) -> tuple[float, float]:
    ajuste_min, ajuste_max = AJUSTE_OBJETIVO[objetivo]
    # para emagrecimento os ajustes são negativos, então min/max se invertem
    valores = sorted([get * (1 + ajuste_min), get * (1 + ajuste_max)])
    return round(valores[0], 0), round(valores[1], 0)


# ---------------------------------------------------------------------------
# Geração de documentos
# ---------------------------------------------------------------------------

def markdown_para_docx(titulo: str, conteudo_md: str, rodape: str = "") -> bytes:
    """Converte um texto em Markdown simples (títulos ##, listas -, texto) em um
    arquivo .docx formatado, retornando os bytes prontos para download."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(titulo, level=0)

    for linha in conteudo_md.splitlines():
        linha = linha.rstrip()
        if not linha:
            continue
        if linha.startswith("### "):
            doc.add_heading(linha[4:], level=3)
        elif linha.startswith("## "):
            doc.add_heading(linha[3:], level=2)
        elif linha.startswith("# "):
            doc.add_heading(linha[2:], level=1)
        elif linha.startswith(("- ", "* ")):
            doc.add_paragraph(linha[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", linha):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", linha), style="List Number")
        else:
            doc.add_paragraph(linha)

    if rodape:
        doc.add_paragraph()
        p = doc.add_paragraph(rodape)
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def markdown_para_pdf(titulo: str, conteudo_md: str, rodape: str = "") -> bytes:
    """Converte um texto em Markdown simples em um PDF pronto para impressão."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, _limpar_para_pdf(titulo))
    pdf.ln(2)

    for linha in conteudo_md.splitlines():
        linha = linha.rstrip()
        if not linha:
            pdf.ln(2)
            continue
        texto = _limpar_para_pdf(linha)
        pdf.set_x(pdf.l_margin)
        if linha.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, texto[3:])
        elif linha.startswith("# "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 8, texto[2:])
        elif linha.startswith(("- ", "* ")):
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 7, f"-  {texto[2:]}")
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 7, texto)

    if rodape:
        pdf.ln(4)
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "I", 8)
        pdf.multi_cell(0, 5, _limpar_para_pdf(rodape))

    saida = pdf.output(dest="S")
    return bytes(saida)


def _limpar_para_pdf(texto: str) -> str:
    """fpdf2 com fontes core (Helvetica) só suporta latin-1; troca emojis/caracteres
    fora desse conjunto por uma aproximação segura para não quebrar a geração."""
    return texto.encode("latin-1", errors="replace").decode("latin-1")


def nome_arquivo(prefixo: str, paciente: str, extensao: str) -> str:
    paciente_slug = re.sub(r"[^a-zA-Z0-9]+", "-", paciente.strip().lower()).strip("-") or "paciente"
    data = datetime.now().strftime("%Y-%m-%d")
    return f"{prefixo}-{paciente_slug}-{data}.{extensao}"
